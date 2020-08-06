import docx
import os #用于获取目标文件所在路径
import win32com
from win32com.client import Dispatch
import win32com.client
import pandas as pd
import numpy as np

def info_update(doc,old_info, new_info):
    '''此函数用于批量替换合同中需要替换的信息
    doc:文件
    old_info和new_info：原文字和需要替换的新文字
    '''
    #读取段落中的所有run，找到需替换的信息进行替换
    for para in doc.paragraphs: #
        for run in para.runs:
            run.text = run.text.replace(old_info, new_info) #替换信息
    #读取表格中的所有单元格，找到需替换的信息进行替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(old_info, new_info) #替换信息

def info_update2(doc,old_info, new_info):
    wordApp = Dispatch('Word.Application')  # 打开word应用程序
    wordApp.Visible = 0  # 后台运行,不显示
    wordApp.DisplayAlerts = 0  # 不警告
    doc = wordApp.Documents.Open(doc, Encoding='gbk')
    # print('段落数',doc.Paragraphs.count)
    # wordApp.Selection.Find.ClearFormatting()
    # wordApp.Selection.Find.Replacement.ClearFormatting()
    wordApp.Selection.Find.ClearFormatting()
    wordApp.Selection.Find.Replacement.ClearFormatting()
    wordApp.Selection.Find.Execute(old_info, False, False, False, False, False, True, 1, True, new_info, 2)
    doc.SaveAs(r'C:\Users\yongjiangao\QA_EVR\替换结果\700-014621-000.docx')
    doc.Close()
# for file in files:
#     doc = docx.Document(file)
#     info_update(doc,"商贸", "贸易")
#     doc.save("data/替换结果/{}".format(file.split("/")[-1]))

if __name__ == '__main__':
    path="C:/Users/yongjiangao/QA_EVR/" # 文件夹路径
    files=[]
    for file in os.listdir(path):
        if file.endswith(".docx"): #排除文件夹内的其它干扰文件，只获取word文件
            files.append(path+file) 
    print(files)

    # for file in files:
    #     doc =docx.Document(file)


    #     # info_update(doc,"MDN", "700-014621-0000")
    #     # info_update(doc,"SPF","700-014621-0000 Test plan v1.0")
    #     # info_update(doc,"PRN","700-014621-0000(QA)Test2A")
    #     # info_update(doc,"RV","2.0")
    #     # info_update(doc,"SPV","1.0")
    #     # info_update(doc,"CS","XX")
    #     # info_update(doc,"SS","650")
    #     # doc.save(r"C:\Users\yongjiangao\QA_EVR\替换结果\700-014621-000.docx")

    for file in files:
        info_update2(file,"MDN", "700-014621-0000")