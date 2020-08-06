import docx
import os #用于获取目标文件所在路径
import win32com
from win32com.client import Dispatch
import win32com.client
import pandas as pd
import numpy as np


# def info_update(doc,old_info, new_info):
#     '''此函数用于批量替换合同中需要替换的信息
#     doc:文件
#     old_info和new_info：原文字和需要替换的新文字
#     '''
#     #读取段落中的所有run，找到需替换的信息进行替换
#     for para in doc.paragraphs: #
#         for run in para.runs:
#             run.text = run.text.replace(old_info, new_info) #替换信息
#     #读取表格中的所有单元格，找到需替换的信息进行替换
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 cell.text = cell.text.replace(old_info, new_info) #替换信息

# def info_update2(doc,old_info, new_info):
#     wordApp = Dispatch('Word.Application')  # 打开word应用程序
#     wordApp.Visible = 0  # 后台运行,不显示
#     wordApp.DisplayAlerts = 0  # 不警告
#     doc = wordApp.Documents.Open(doc, Encoding='gbk')
#     # print('段落数',doc.Paragraphs.count)
#     # wordApp.Selection.Find.ClearFormatting()
#     # wordApp.Selection.Find.Replacement.ClearFormatting()
#     wordApp.Selection.Find.ClearFormatting()
#     wordApp.Selection.Find.Replacement.ClearFormatting()
#     wordApp.Selection.Find.Execute(old_info, False, False, False, False, False, True, 1, True, new_info, 2)
#     doc.SaveAs(r'C:\Users\yongjiangao\QA_EVR\替换结果\700-014621-000.docx')
#     doc.Close()
# for file in files:
#     doc = docx.Document(file)
#     info_update(doc,"商贸", "贸易")
#     doc.save("data/替换结果/{}".format(file.split("/")[-1]))

if __name__ == '__main__':
    EvrFilePath =r"C:\Users\yongjiangao\ProgramEvr_02.xlsx"
    path = r"C:\Users\yongjiangao\QA_EVR\QA.docx"
    EvaData=pd.DataFrame(pd.read_excel(EvrFilePath,sheet_name="ZJ2 (4)"))
    wordApp = win32com.client.Dispatch('Word.Application')  # 打开word应用程序
    wordApp.Visible = 0  # 后台运行,不显示
    wordApp.DisplayAlerts = 0  # 不警告
    doc = wordApp.Documents.Open(path, Encoding='gbk')
    wordApp.Selection.Find.ClearFormatting()
    wordApp.Selection.Find.Replacement.ClearFormatting()
    for row in range(EvaData.shape[0]):
        # Pathtemp=r("C:\Users\yongjiangao\QA_EVR\替换结果")
        doc = wordApp.Documents.Open(path, Encoding='gbk')
        wordApp.Selection.Find.ClearFormatting()
        wordApp.Selection.Find.Replacement.ClearFormatting()
        newPath="C:\\Users\\yongjiangao\\QA_EVR\\替换结果\\"+str(EvaData['Model description'][row])+" "+str(EvaData['Flow2'][row])+" QA EVR.docx"
        print(EvaData['Model description'][row])
        print(EvaData['Flow2'][row])
        print(EvaData['StanderProgramName'][row])
        print(EvaData['SpecPath'][row])
        print(EvaData['TestProgram Checksum'][row])
        print(EvaData['Test Program Rev'][row])
        # xlSheet= xlBook.Worksheets('sheet1')
        # xlSheet.Cells(7,2).Value="Application Name:  "+str(EvaData['Model description'][row])+" QA Program"
        # xlSheet.Cells(9,2).Value="Application Name:  "+str(EvaData['StanderProgramName'][row])+" Rev:"+str(EvaData['Test Program Rev'][row])
        # xlSheet.Cells(11,2).Value="Designed Use :  "+str(EvaData['Model description'][row])
        # xlSheet.Cells(20,4).Value=EvaData['TestProgram Checksum'][row]
        # xlSheet.Cells(21,2).Value="Location of the test spec :"+str(EvaData['SpecPath'][row])
        # xlBook.SaveAs(newPath)
        wordApp.Selection.Find.Execute("MDN", False, False, False, False, False, True, 1, True, EvaData['Model description'][row], 2)
        wordApp.Selection.Find.Execute("PRN", False, False, False, False, False, True, 1, True, EvaData['StanderProgramName'][row], 2)
        wordApp.Selection.Find.Execute("RV", False, False, False, False, False, True, 1, True, EvaData['Test Program Rev'][row], 2)
        wordApp.Selection.Find.Execute("SPF", False, False, False, False, False, True, 1, True, EvaData['SpecName'][row], 2)
        wordApp.Selection.Find.Execute("SPV", False, False, False, False, False, True, 1, True, EvaData['SpecRev'][row], 2)
        wordApp.Selection.Find.Execute("CS", False, False, False, False, False, True, 1, True, EvaData['TestProgram Checksum'][row], 2)
        wordApp.Selection.Find.Execute("SS", False, False, False, False, False, True, 1, True, EvaData['TestTime'][row], 2)
        doc.SaveAs(newPath)
        print("---------")
        doc.Close()